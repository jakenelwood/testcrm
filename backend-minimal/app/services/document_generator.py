import os
import json
import datetime
import re
import logging
from docx import Document as DocxDocument
from pathlib import Path
import uuid
import subprocess
from sqlalchemy.orm import Session
from app.models.quote import Quote
from app.models.client import Client
import traceback

# Set up logging
logger = logging.getLogger(__name__)
handler = logging.StreamHandler()
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)
logger.addHandler(handler)
logger.setLevel(logging.INFO)

# Create document storage directory
DOCUMENT_DIR = Path("./documents")
DOCUMENT_DIR.mkdir(exist_ok=True)

# Template directory
TEMPLATE_DIR = Path("./app/templates")
if not TEMPLATE_DIR.exists():
    # Check if templates are in the root directory
    if Path("./auto-request-form.docx").exists():
        TEMPLATE_DIR = Path(".")
    else:
        logger.warning("Template directory not found, creating empty directory")
        TEMPLATE_DIR.mkdir(exist_ok=True)

# Template filenames
AUTO_TEMPLATE = "auto-request-form.docx"
HOME_TEMPLATE = "home-quote-request-form.docx"
SPECIALTY_TEMPLATE = "specialty-quote-request-form.docx"

def generate_quote_document(quote: Quote, file_type: str = "docx"):
    """
    Generate a document from a quote using templates and placeholder replacement.
    
    Args:
        quote: The quote object
        file_type: Either 'docx' or 'pdf'
        
    Returns:
        Tuple of (filename, file_path)
    
    Raises:
        ValueError: If an unsupported file type is provided
        Exception: For any errors during document generation
    """
    try:
        logger.info(f"Starting document generation for quote {quote.id}, type: {file_type}")
        
        # Get client data
        client = quote.client
        if not client:
            logger.error(f"No client data found for quote {quote.id}")
            raise ValueError(f"No client data found for quote {quote.id}")
        
        # Create a unique filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        base_filename = f"quote_{quote.id}_{client.name.replace(' ', '_')}_{timestamp}_{unique_id}"
        
        logger.info(f"Generating {file_type} document with base filename: {base_filename}")
        
        if file_type.lower() == "docx":
            result = _generate_docx(quote, client, base_filename)
            logger.info(f"DOCX generation completed: {result[0]}")
            return result
        elif file_type.lower() == "pdf":
            result = _generate_pdf(quote, client, base_filename)
            logger.info(f"PDF generation completed: {result[0]}")
            return result
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as e:
        logger.error(f"Error generating document: {str(e)}")
        logger.error(traceback.format_exc())
        raise

def _generate_docx(quote: Quote, client: Client, base_filename: str):
    """
    Generate a DOCX document using templates and placeholder replacement.
    """
    try:
        # Determine which templates to use based on quote type
        templates_to_use = []
        
        if quote.has_auto:
            logger.info(f"Adding auto template for quote {quote.id}")
            auto_template_path = TEMPLATE_DIR / AUTO_TEMPLATE
            if auto_template_path.exists():
                templates_to_use.append(auto_template_path)
            else:
                logger.warning(f"Auto template not found at {auto_template_path}")
        
        if quote.has_home:
            logger.info(f"Adding home template for quote {quote.id}")
            home_template_path = TEMPLATE_DIR / HOME_TEMPLATE
            if home_template_path.exists():
                templates_to_use.append(home_template_path)
            else:
                logger.warning(f"Home template not found at {home_template_path}")
        
        if quote.has_specialty:
            logger.info(f"Adding specialty template for quote {quote.id}")
            specialty_template_path = TEMPLATE_DIR / SPECIALTY_TEMPLATE
            if specialty_template_path.exists():
                templates_to_use.append(specialty_template_path)
            else:
                logger.warning(f"Specialty template not found at {specialty_template_path}")
        
        # If no specific insurance types were selected, use all templates
        if not templates_to_use:
            logger.warning(f"No insurance types specified for quote {quote.id}, using all templates")
            for template_name in [AUTO_TEMPLATE, HOME_TEMPLATE, SPECIALTY_TEMPLATE]:
                template_path = TEMPLATE_DIR / template_name
                if template_path.exists():
                    templates_to_use.append(template_path)
                else:
                    logger.warning(f"Template not found: {template_path}")
        
        # If still no templates found, raise an error
        if not templates_to_use:
            logger.error("No templates found")
            raise ValueError("No templates found for document generation")
        
        # Create a data dictionary for placeholder replacement
        data = _create_data_dictionary(quote, client)
        
        # Create a new document
        doc = None
        
        # Process each template
        for template_path in templates_to_use:
            logger.info(f"Processing template: {template_path}")
            template_doc = DocxDocument(str(template_path))
            
            # Process placeholders in the template
            _process_docx_placeholders(template_doc, data)
            
            # For the first template, use it as the base document
            if doc is None:
                doc = template_doc
            else:
                # For subsequent templates, append content to the base document
                for element in template_doc.element.body:
                    doc.element.body.append(element)
        
        # Save the document
        filename = f"{base_filename}.docx"
        file_path = str(DOCUMENT_DIR / filename)
        logger.info(f"Saving DOCX to: {file_path}")
        doc.save(file_path)
        
        logger.info(f"Document saved: {file_path}")
        return filename, file_path
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        logger.error(traceback.format_exc())
        raise

def _generate_pdf(quote: Quote, client: Client, base_filename: str):
    """
    Generate a PDF document by first creating a DOCX and then converting it to PDF.
    """
    try:
        logger.info(f"Generating PDF for quote {quote.id}")
        
        # First generate the DOCX file
        docx_filename, docx_path = _generate_docx(quote, client, base_filename)
        
        # Define PDF path
        pdf_filename = f"{base_filename}.pdf"
        pdf_path = str(DOCUMENT_DIR / pdf_filename)
        
        # Convert DOCX to PDF using LibreOffice (platform-independent approach)
        logger.info(f"Converting DOCX to PDF: {docx_path} -> {pdf_path}")
        
        # Try different PDF conversion methods
        conversion_methods = [
            _convert_with_libreoffice,
            _convert_with_unoconv,
            _convert_with_fallback
        ]
        
        # Try each conversion method in order
        for method in conversion_methods:
            try:
                logger.info(f"Attempting PDF conversion with {method.__name__}")
                success = method(docx_path, pdf_path)
                if success:
                    logger.info(f"PDF conversion successful using {method.__name__}")
                    return pdf_filename, pdf_path
                else:
                    logger.warning(f"Conversion method {method.__name__} failed")
            except Exception as e:
                logger.warning(f"Conversion method {method.__name__} failed: {str(e)}")
                # Continue to next method
        
        # If we get here, all conversion methods failed
        logger.error("All PDF conversion methods failed. Creating placeholder PDF.")
        with open(docx_path, 'rb') as src_file:
            with open(pdf_path, 'wb') as dst_file:
                dst_file.write(src_file.read())
        
        logger.warning(f"Created placeholder PDF (actually DOCX): {pdf_path}")
        
        return pdf_filename, pdf_path
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        logger.error(traceback.format_exc())
        raise

def _convert_with_libreoffice(docx_path: str, pdf_path: str) -> bool:
    """
    Convert DOCX to PDF using LibreOffice
    
    Returns:
        True if conversion successful, False otherwise
    """
    try:
        # Try using LibreOffice for conversion
        conversion_cmd = [
            'libreoffice', 
            '--headless', 
            '--convert-to', 
            'pdf', 
            '--outdir', 
            str(DOCUMENT_DIR),
            docx_path
        ]
        
        # Execute the command
        process = subprocess.run(
            conversion_cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            check=True,
            timeout=60  # 60 second timeout
        )
        
        logger.info(f"LibreOffice conversion output: {process.stdout}")
        
        # LibreOffice creates the PDF with the same base name, so we need to rename it
        libreoffice_pdf = str(DOCUMENT_DIR / f"{os.path.basename(docx_path).replace('.docx', '.pdf')}")
        if os.path.exists(libreoffice_pdf) and libreoffice_pdf != pdf_path:
            os.rename(libreoffice_pdf, pdf_path)
            
        if os.path.exists(pdf_path):
            logger.info(f"PDF created successfully at: {pdf_path}")
            return True
        else:
            logger.warning(f"PDF file not created by LibreOffice at expected location: {pdf_path}")
            return False
            
    except Exception as e:
        logger.warning(f"Error using LibreOffice for conversion: {str(e)}")
        return False

def _convert_with_unoconv(docx_path: str, pdf_path: str) -> bool:
    """
    Convert DOCX to PDF using unoconv
    
    Returns:
        True if conversion successful, False otherwise
    """
    try:
        # Try using unoconv for conversion
        conversion_cmd = [
            'unoconv',
            '-f', 
            'pdf', 
            '-o',
            pdf_path,
            docx_path
        ]
        
        # Execute the command
        process = subprocess.run(
            conversion_cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            check=True,
            timeout=30  # 30 second timeout
        )
        
        logger.info(f"Unoconv conversion output: {process.stdout}")
        
        # Check if conversion was successful
        if os.path.exists(pdf_path):
            return True
        
        return False
    except Exception as e:
        logger.warning(f"Unoconv conversion failed: {str(e)}")
        return False

def _convert_with_fallback(docx_path: str, pdf_path: str) -> bool:
    """
    Fallback method that creates a placeholder PDF
    (Just copies the DOCX file)
    
    Returns:
        True if successful, False otherwise
    """
    try:
        logger.warning("Using fallback PDF conversion (copying DOCX file)")
        with open(docx_path, 'rb') as src_file:
            with open(pdf_path, 'wb') as dst_file:
                dst_file.write(src_file.read())
        
        logger.warning(f"Created placeholder PDF (actually DOCX): {pdf_path}")
        return True
    except Exception as e:
        logger.error(f"Fallback conversion failed: {str(e)}")
        return False

def _create_data_dictionary(quote: Quote, client: Client):
    """
    Create a dictionary mapping placeholders to values from the quote and client data.
    """
    try:
        # Current date in MM/DD/YYYY format
        current_date = datetime.datetime.now().strftime("%m/%d/%Y")
        
        # Start with client data
        data = {
            # Basic client information
            "pniname": client.name or "",
            "pniphone": client.phone_number or "",
            "pniemail": client.email or "",
            "pniaddr": client.address or "",
            "pni-mailingaddr": client.mailing_address or "",
            "pni-prioraddr": client.prior_address or "",
            
            # Extended client information
            "pnig": client.gender or "",
            "pnims": client.marital_status or "",
            "pnidob": _format_date(client.date_of_birth) or "",
            "pniedocc": client.education_occupation or "",
            "pnidln": client.drivers_license or "",
            "pnidls": client.license_state or "",
            "pnissn": client.ssn or "",
            "pnirel": client.relation or "",
            
            # Referral info
            "referred-by": client.referred_by or "",
            
            # Quote information
            "current-date": current_date,
            "effective-date": _format_date(quote.effective_date) or "",
        }
        
        # Add auto data if present
        if quote.has_auto and quote.auto_data:
            logger.info(f"Processing auto data for quote {quote.id}")
            _add_auto_data_to_dictionary(data, quote.auto_data)
        
        # Add home data if present
        if quote.has_home and quote.home_data:
            logger.info(f"Processing home data for quote {quote.id}")
            _add_home_data_to_dictionary(data, quote.home_data)
        
        # Add specialty data if present
        if quote.has_specialty and quote.specialty_data:
            logger.info(f"Processing specialty data for quote {quote.id}")
            _add_specialty_data_to_dictionary(data, quote.specialty_data)
        
        # Add additional insured data if present in auto, home, or specialty data
        _add_additional_insured_data(data, quote)
        
        # Default empty string for any missing placeholders
        # This ensures the template doesn't fail if a placeholder is present but not in our data
        logger.info(f"Created data dictionary with {len(data)} entries")
        return data
    except Exception as e:
        logger.error(f"Error creating data dictionary: {str(e)}")
        logger.error(traceback.format_exc())
        raise

def _add_auto_data_to_dictionary(data, auto_data):
    """
    Add auto insurance data to the data dictionary
    """
    # Map basic auto fields
    auto_field_mapping = {
        "current_carrier": "a-current-carrier",
        "months_with_carrier": "a-mos-current-carrier", 
        "current_limits": "a-climits",
        "quoting_limits": "a-qlimits",
        "expiration_date": "a-exp-dt",
        "premium": "aprem",
        "additional_notes": "auto-additional-notes"
    }
    
    # Process basic auto fields
    for api_key, placeholder_key in auto_field_mapping.items():
        if api_key in auto_data:
            value = auto_data[api_key]
            # Format date fields
            if api_key == "expiration_date":
                value = _format_date(value)
            data[placeholder_key] = value if value is not None else ""
    
    # Process driver data 
    if "drivers" in auto_data and isinstance(auto_data["drivers"], list):
        for i, driver in enumerate(auto_data["drivers"], 1):
            if i > 8:  # We only support up to 8 drivers in the template
                logger.warning(f"More than 8 drivers found, ignoring drivers beyond the 8th")
                break
                
            # Driver field mapping
            driver_field_mapping = {
                "firstName": f"d{i}first",
                "lastName": f"d{i}last",
                "dateOfBirth": f"d{i}dob",
                "licenseNumber": f"d{i}license",
                "licenseState": f"d{i}state",
                "gender": f"d{i}gender",
                "maritalStatus": f"d{i}marital"
            }
            
            # Add all driver fields to data dictionary
            for api_key, placeholder_key in driver_field_mapping.items():
                if api_key in driver:
                    value = driver[api_key]
                    # Format date fields
                    if api_key == "dateOfBirth":
                        value = _format_date(value)
                    data[placeholder_key] = value if value is not None else ""
                else:
                    # Set default for missing fields
                    data[placeholder_key] = ""
    
    # Process vehicle data
    if "vehicles" in auto_data and isinstance(auto_data["vehicles"], list):
        for i, vehicle in enumerate(auto_data["vehicles"], 1):
            if i > 8:  # We only support up to 8 vehicles in the template
                logger.warning(f"More than 8 vehicles found, ignoring vehicles beyond the 8th")
                break
                
            # Vehicle field mapping (API field name to placeholder field name)
            vehicle_field_mapping = {
                "year": f"v{i}yr",
                "make": f"v{i}make",
                "model": f"v{i}model",
                "vin": f"v{i}vin",
                "usage": f"v{i}usage",
                "mileage": f"v{i}mi",
                "driver": f"v{i}-driver",
                "comprehensive": f"v{i}comp",
                "collision": f"v{i}coll",
                "glass": f"v{i}glass",
                "towing": f"v{i}tow",
                "rental_reimbursement": f"v{i}rr",
                "financing": f"v{i}fin",
                "gap": f"v{i}gap"
            }
            
            # Add all vehicle fields to data dictionary
            for api_key, placeholder_key in vehicle_field_mapping.items():
                if api_key in vehicle:
                    value = vehicle[api_key]
                    # Convert boolean values to "Yes"/"No" strings
                    if isinstance(value, bool):
                        value = "Yes" if value else "No"
                    data[placeholder_key] = value if value is not None else ""
                else:
                    # Set default for missing fields
                    data[placeholder_key] = ""

def _add_home_data_to_dictionary(data, home_data):
    """
    Add home insurance data to the data dictionary
    """
    # Process each field in home_data
    for key, value in home_data.items():
        # Convert home_data keys to placeholder format
        placeholder_key = key.replace("_", "-")
        if not placeholder_key.startswith("h"):
            placeholder_key = f"h{placeholder_key}"
        
        # Handle special cases for home data fields
        if key == "expiration_date":
            value = _format_date(value)
        
        # Convert boolean values to "Yes"/"No" strings
        if isinstance(value, bool):
            value = "Yes" if value else "No"
            
        data[placeholder_key] = value if value is not None else ""

def _add_specialty_data_to_dictionary(data, specialty_data):
    """
    Add specialty insurance data to the data dictionary
    """
    # Handle specialty items (assuming it's in a nested structure)
    if "items" in specialty_data and isinstance(specialty_data["items"], list):
        for i, item in enumerate(specialty_data["items"], 1):
            if i > 8:  # We only support up to 8 specialty items in the template
                logger.warning(f"More than 8 specialty items found, ignoring items beyond the 8th")
                break
                
            # Specialty field mapping
            for skey, svalue in item.items():
                # Convert specialty keys to placeholder format
                placeholder_key = f"sp{i}-{skey.replace('_', '-')}" if not skey.startswith("type") else f"sp{i}-type-toy"
                
                # Convert boolean values to "Yes"/"No" strings
                if isinstance(svalue, bool):
                    svalue = "Yes" if svalue else "No"
                    
                data[placeholder_key] = svalue if svalue is not None else ""
    
    # Add additional info if present
    if "additional_info" in specialty_data:
        data["sp-additional-info"] = specialty_data["additional_info"] or ""

def _add_additional_insured_data(data, quote):
    """
    Add additional insured data to the data dictionary
    """
    for insurance_type in ["auto_data", "home_data", "specialty_data"]:
        insurance_data = getattr(quote, insurance_type, None)
        if not insurance_data or not isinstance(insurance_data, dict):
            continue
            
        if "additional_insureds" in insurance_data:
            additional_insureds = insurance_data["additional_insureds"]
            if isinstance(additional_insureds, list):
                for i, insured in enumerate(additional_insureds, 2):  # Start at 2 because 1 is the primary insured
                    if i > 8:  # We only support up to 8 additional insureds in the template
                        logger.warning(f"More than 8 additional insureds found, ignoring beyond the 8th")
                        break
                        
                    # Additional insured field mapping
                    ai_field_mapping = {
                        "name": f"ai{i}name",
                        "relation": f"ai{i}rel",
                        "gender": f"ai{i}g",
                        "marital_status": f"ai{i}ms",
                        "date_of_birth": f"ai{i}dob",
                        "education_occupation": f"ai{i}edocc",
                        "drivers_license": f"ai{i}dln",
                        "license_state": f"ai{i}dls",
                        "ssn": f"ai{i}ssn"
                    }
                    
                    # Add all additional insured fields
                    for api_key, placeholder_key in ai_field_mapping.items():
                        if api_key in insured:
                            value = insured[api_key]
                            # Format date fields
                            if api_key == "date_of_birth":
                                value = _format_date(value)
                            data[placeholder_key] = value if value is not None else ""
                        else:
                            # Set default for missing fields
                            data[placeholder_key] = ""

def _process_docx_placeholders(doc, data):
    """
    Process all placeholders in a DOCX document.
    
    This includes:
    - Text in paragraphs
    - Text in tables
    - Text in headers/footers
    """
    try:
        # Define the regex pattern for placeholders: {{placeholder}}
        pattern = r'{{(.*?)}}'
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            _replace_text_in_paragraph(paragraph, pattern, data)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_text_in_paragraph(paragraph, pattern, data)
        
        # Replace in headers and footers
        for section in doc.sections:
            for header in [section.header, section.footer]:
                if header:
                    for paragraph in header.paragraphs:
                        _replace_text_in_paragraph(paragraph, pattern, data)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    _replace_text_in_paragraph(paragraph, pattern, data)
    except Exception as e:
        logger.error(f"Error processing placeholders: {str(e)}")
        logger.error(traceback.format_exc())
        raise

def _replace_text_in_paragraph(paragraph, pattern, data):
    """
    Replace placeholders in a paragraph with values from the data dictionary.
    """
    if not paragraph.text:
        return
    
    # Find all matches in the paragraph text
    matches = re.findall(pattern, paragraph.text)
    
    # If no matches found, return
    if not matches:
        return
    
    # For each match, replace the placeholder with the corresponding value
    for placeholder in matches:
        placeholder_value = data.get(placeholder, "")
        if placeholder not in data:
            logger.warning(f"Placeholder '{placeholder}' not found in data dictionary")
        paragraph.text = paragraph.text.replace(f"{{{{{placeholder}}}}}", str(placeholder_value))

def _format_date(date_value):
    """
    Format a date value to MM/DD/YYYY string.
    
    Args:
        date_value: Date as string, datetime object, or any other format
        
    Returns:
        Formatted date string or empty string if invalid
    """
    if not date_value:
        return ""
        
    try:
        # If it's already a string in correct format, return it
        if isinstance(date_value, str):
            try:
                # Try to parse as date and reformat
                parsed_date = datetime.datetime.strptime(date_value, "%Y-%m-%d")
                return parsed_date.strftime("%m/%d/%Y")
            except ValueError:
                # If can't parse, return as is
                return date_value
                
        # If it's a datetime object, format it
        if isinstance(date_value, (datetime.datetime, datetime.date)):
            return date_value.strftime("%m/%d/%Y")
            
        # If it's something else, convert to string
        return str(date_value)
    except Exception as e:
        logger.error(f"Error formatting date '{date_value}': {str(e)}")
        return ""
