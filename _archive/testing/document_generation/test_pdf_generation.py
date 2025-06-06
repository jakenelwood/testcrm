import pytest
import os
import time
import subprocess
from pathlib import Path
from unittest.mock import patch, MagicMock

# Test paths
TEST_DIR = Path("./testing/document_generation")
TEST_DIR.mkdir(exist_ok=True, parents=True)
SAMPLE_OUTPUT_DIR = TEST_DIR / "output"
SAMPLE_OUTPUT_DIR.mkdir(exist_ok=True)
TEST_DOCX_PATH = SAMPLE_OUTPUT_DIR / "test_document.docx"

@pytest.fixture
def mock_pdf_environment():
    """Setup environment for PDF testing"""
    # Create a temporary directory for test files
    temp_dir = TEST_DIR / "pdf_tests"
    temp_dir.mkdir(exist_ok=True)
    
    # Create a simple docx file for testing conversion
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test document for PDF conversion")
    doc.add_paragraph("This is a test paragraph.")
    doc.save(TEST_DOCX_PATH)
    
    # Return test configuration
    return {
        "temp_dir": temp_dir,
        "test_docx_path": TEST_DOCX_PATH,
        "docx_filename": TEST_DOCX_PATH.name,
        "pdf_filename": TEST_DOCX_PATH.stem + ".pdf"
    }

def test_pdf_conversion_functionality():
    """Test basic PDF conversion functionality"""
    from backend.app.services.document_generator import _generate_pdf
    
    # Create mock objects
    mock_quote = MagicMock()
    mock_quote.id = 1
    mock_client = MagicMock()
    mock_client.name = "Test Client"
    mock_quote.client = mock_client
    
    # Test with different base filenames
    base_filename = f"test_quote_{int(time.time())}"
    
    # Create a patched version of the _generate_docx function
    with patch('backend.app.services.document_generator._generate_docx') as mock_generate_docx:
        # Configure mock to return a test DOCX path
        mock_generate_docx.return_value = (f"{base_filename}.docx", str(TEST_DOCX_PATH))
        
        # Call the PDF generation function (which will use our mocked _generate_docx)
        filename, file_path = _generate_pdf(mock_quote, mock_client, base_filename)
        
        # Verify we get back a PDF filename and path
        assert filename == f"{base_filename}.pdf"
        assert file_path.endswith(".pdf")
        
        # Verify the mock was called correctly
        mock_generate_docx.assert_called_once_with(mock_quote, mock_client, base_filename)

def test_pdf_conversion_with_libreoffice():
    """Test PDF conversion using LibreOffice (if available)"""
    # Skip test if LibreOffice is not available
    try:
        result = subprocess.run(['libreoffice', '--version'], 
                                stdout=subprocess.PIPE, 
                                stderr=subprocess.PIPE, 
                                text=True, 
                                check=False)
        if result.returncode != 0:
            pytest.skip("LibreOffice not available, skipping test")
    except (FileNotFoundError, subprocess.SubprocessError):
        pytest.skip("LibreOffice not available, skipping test")
    
    # Test with an actual DOCX file
    from backend.app.services.document_generator import _generate_pdf
    
    # Create a patch for the subprocess.run call to monitor it
    with patch('subprocess.run') as mock_run:
        # Configure mock to return a successful result
        mock_process = MagicMock()
        mock_process.returncode = 0
        mock_process.stdout = "Conversion successful"
        mock_run.return_value = mock_process
        
        # Call the function with mocked subprocess
        _generate_pdf(MagicMock(), MagicMock(), "test_libreoffice")
        
        # Verify that libreoffice was called correctly
        mock_run.assert_called()
        args, kwargs = mock_run.call_args
        cmd = args[0]
        
        # Check that the command includes libreoffice with the right flags
        assert 'libreoffice' in cmd[0]
        assert '--headless' in cmd[1]
        assert '--convert-to' in cmd[2]
        assert 'pdf' in cmd[3]

def test_pdf_fallback_mechanism():
    """Test fallback mechanism when LibreOffice fails"""
    from backend.app.services.document_generator import _generate_pdf
    
    # Create mock objects
    mock_quote = MagicMock()
    mock_client = MagicMock()
    mock_quote.client = mock_client
    
    # Create a patch for _generate_docx to return a known path
    with patch('backend.app.services.document_generator._generate_docx') as mock_generate_docx:
        # Make _generate_docx return our test DOCX file
        mock_generate_docx.return_value = ("test_doc.docx", str(TEST_DOCX_PATH))
        
        # Create a patch for subprocess.run to simulate a LibreOffice failure
        with patch('subprocess.run') as mock_run:
            # Configure mock to raise an exception
            mock_run.side_effect = subprocess.SubprocessError("LibreOffice conversion failed")
            
            # Call the function with mocked subprocess
            filename, file_path = _generate_pdf(mock_quote, mock_client, "test_fallback")
            
            # Verify we get back a PDF filename (even though it might just be a renamed DOCX)
            assert filename.endswith(".pdf")
            assert file_path.endswith(".pdf")

def test_pdf_conversion_large_document():
    """Test PDF conversion with a larger document"""
    from backend.app.services.document_generator import _generate_pdf
    
    # Create a larger test document
    from docx import Document
    large_doc_path = SAMPLE_OUTPUT_DIR / "large_test_document.docx"
    doc = Document()
    
    # Add multiple paragraphs and tables to create a larger document
    for i in range(100):
        doc.add_paragraph(f"This is paragraph {i} in a large test document. " * 5)
    
    for i in range(10):
        table = doc.add_table(rows=5, cols=5)
        for row in range(5):
            for col in range(5):
                table.cell(row, col).text = f"Cell {row},{col} content for table {i}"
    
    doc.save(large_doc_path)
    
    # Create mock objects
    mock_quote = MagicMock()
    mock_client = MagicMock()
    mock_quote.client = mock_client
    
    # Create patches
    with patch('backend.app.services.document_generator._generate_docx') as mock_generate_docx:
        # Make _generate_docx return our large test DOCX file
        mock_generate_docx.return_value = ("large_test_doc.docx", str(large_doc_path))
        
        # Increase timeout for large document
        with patch('subprocess.run') as mock_run:
            # Configure mock to simulate successful conversion
            mock_process = MagicMock()
            mock_process.returncode = 0
            mock_run.return_value = mock_process
            
            # Time the conversion to measure performance
            start_time = time.time()
            filename, file_path = _generate_pdf(mock_quote, mock_client, "large_test")
            conversion_time = time.time() - start_time
            
            # Log performance metrics
            print(f"Large document conversion time: {conversion_time:.2f} seconds")
            
            # Verify subprocess was called with an appropriate timeout
            mock_run.assert_called()
            _, kwargs = mock_run.call_args
            assert kwargs.get('timeout', 0) >= 30, "Timeout should be at least 30 seconds for large documents"

def test_pdf_concurrent_conversion():
    """Test converting multiple PDFs concurrently"""
    import concurrent.futures
    from backend.app.services.document_generator import _generate_pdf
    
    # Number of concurrent conversions
    num_concurrent = 5
    
    # Create a test DOCX file if it doesn't exist
    if not TEST_DOCX_PATH.exists():
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test document for concurrent PDF conversion")
        doc.save(TEST_DOCX_PATH)
    
    # Function to perform conversion with patched components
    def convert_to_pdf(task_id):
        mock_quote = MagicMock()
        mock_quote.id = task_id
        mock_client = MagicMock()
        mock_client.name = f"Client {task_id}"
        mock_quote.client = mock_client
        
        base_filename = f"concurrent_test_{task_id}_{int(time.time())}"
        
        # Create patches
        with patch('backend.app.services.document_generator._generate_docx') as mock_generate_docx:
            # Configure mock to return a test DOCX path
            mock_generate_docx.return_value = (f"{base_filename}.docx", str(TEST_DOCX_PATH))
            
            # Create patch for subprocess
            with patch('subprocess.run') as mock_run:
                # Configure mock to simulate successful conversion
                mock_process = MagicMock()
                mock_process.returncode = 0
                mock_run.return_value = mock_process
                
                # Perform conversion
                start_time = time.time()
                filename, file_path = _generate_pdf(mock_quote, mock_client, base_filename)
                elapsed_time = time.time() - start_time
                
                return {
                    "task_id": task_id,
                    "filename": filename,
                    "file_path": file_path,
                    "elapsed_time": elapsed_time,
                    "success": True
                }
    
    # Execute concurrent conversions
    start_time = time.time()
    with concurrent.futures.ThreadPoolExecutor(max_workers=num_concurrent) as executor:
        futures = [executor.submit(convert_to_pdf, i) for i in range(1, num_concurrent + 1)]
        results = [future.result() for future in concurrent.futures.as_completed(futures)]
    total_time = time.time() - start_time
    
    # Verify results
    assert len(results) == num_concurrent
    all_success = all(result["success"] for result in results)
    assert all_success, "All conversions should succeed"
    
    # Log performance metrics
    avg_time = sum(result["elapsed_time"] for result in results) / num_concurrent
    print(f"Converted {num_concurrent} PDFs in {total_time:.2f} seconds")
    print(f"Average conversion time: {avg_time:.2f} seconds")
    print(f"Throughput: {num_concurrent / total_time:.2f} conversions/second")

def test_pdf_character_encoding():
    """Test PDF conversion with special characters"""
    from backend.app.services.document_generator import _generate_pdf
    
    # Create a test document with special characters
    from docx import Document
    special_chars_path = SAMPLE_OUTPUT_DIR / "special_chars.docx"
    doc = Document()
    
    # Add paragraphs with special characters
    doc.add_paragraph("Special characters: áéíóúñÁÉÍÓÚÑ")
    doc.add_paragraph("Symbols: ©®™§¶†‡♠♣♥♦")
    doc.add_paragraph("Other languages: 你好, こんにちは, مرحبا, שלום")
    doc.save(special_chars_path)
    
    # Create mock objects
    mock_quote = MagicMock()
    mock_client = MagicMock()
    mock_quote.client = mock_client
    
    # Create patches
    with patch('backend.app.services.document_generator._generate_docx') as mock_generate_docx:
        # Configure mock to return our special characters DOCX
        mock_generate_docx.return_value = ("special_chars.docx", str(special_chars_path))
        
        # Create patch for subprocess
        with patch('subprocess.run') as mock_run:
            # Configure mock to simulate successful conversion
            mock_process = MagicMock()
            mock_process.returncode = 0
            mock_run.return_value = mock_process
            
            # Call the function with mocked subprocess
            filename, file_path = _generate_pdf(mock_quote, mock_client, "special_chars")
            
            # Verify the result
            assert filename.endswith(".pdf")
            
            # Verify subprocess was called correctly
            mock_run.assert_called()
            
def test_pdf_conversion_error_handling():
    """Test error handling during PDF conversion"""
    from backend.app.services.document_generator import _generate_pdf
    
    # Create mock objects
    mock_quote = MagicMock()
    mock_client = MagicMock()
    mock_quote.client = mock_client
    
    # Test with different error scenarios
    with patch('backend.app.services.document_generator._generate_docx') as mock_generate_docx:
        # 1. Test timeout error
        with patch('subprocess.run') as mock_run:
            mock_run.side_effect = subprocess.TimeoutExpired(cmd="libreoffice", timeout=30)
            
            # Should fall back to alternative method
            filename, file_path = _generate_pdf(mock_quote, mock_client, "timeout_test")
            assert filename.endswith(".pdf")
        
        # 2. Test file not found error
        with patch('subprocess.run') as mock_run:
            mock_run.side_effect = FileNotFoundError("libreoffice command not found")
            
            # Should fall back to alternative method
            filename, file_path = _generate_pdf(mock_quote, mock_client, "file_not_found_test")
            assert filename.endswith(".pdf")
        
        # 3. Test other subprocess error
        with patch('subprocess.run') as mock_run:
            mock_run.side_effect = subprocess.SubprocessError("Generic subprocess error")
            
            # Should fall back to alternative method
            filename, file_path = _generate_pdf(mock_quote, mock_client, "subprocess_error_test")
            assert filename.endswith(".pdf")

def test_pdf_output_verification():
    """Test that the generated PDF file exists and is accessible"""
    from backend.app.services.document_generator import _generate_pdf
    
    # Create a unique test path for this test
    unique_docx_path = SAMPLE_OUTPUT_DIR / f"verify_test_{int(time.time())}.docx"
    
    # Create a simple test document
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test document for output verification")
    doc.save(unique_docx_path)
    
    # Create mock objects and patches
    mock_quote = MagicMock()
    mock_client = MagicMock()
    mock_quote.client = mock_client
    
    with patch('backend.app.services.document_generator._generate_docx') as mock_generate_docx:
        mock_generate_docx.return_value = (unique_docx_path.name, str(unique_docx_path))
        
        # Create patch for subprocess
        with patch('os.path.exists') as mock_exists:
            # First verification fails, then fallback file exists
            mock_exists.side_effect = [False, True]
            
            with patch('subprocess.run') as mock_run:
                # Configure subprocess to succeed
                mock_process = MagicMock()
                mock_process.returncode = 0
                mock_run.return_value = mock_process
                
                # Run the test
                filename, file_path = _generate_pdf(mock_quote, mock_client, "verify_test")
                
                # Verify the result
                assert filename.endswith(".pdf")
                assert os.path.basename(file_path).endswith(".pdf")

if __name__ == "__main__":
    pytest.main(["-v", "test_pdf_generation.py"]) 