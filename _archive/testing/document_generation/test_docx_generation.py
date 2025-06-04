import pytest
import os
import time
from pathlib import Path
from unittest.mock import patch, MagicMock
import docx

# Test paths
TEST_DIR = Path("./testing/document_generation")
TEST_DIR.mkdir(exist_ok=True, parents=True)
SAMPLE_OUTPUT_DIR = TEST_DIR / "output"
SAMPLE_OUTPUT_DIR.mkdir(exist_ok=True)

@pytest.fixture
def mock_quote():
    """Create a comprehensive mock quote with all data types for testing"""
    client = MagicMock()
    client.name = "Test Client"
    client.phone_number = "555-123-4567"
    client.email = "test@example.com"
    client.address = "123 Test St, Test City, TS 12345"
    client.mailing_address = "123 Test St, Test City, TS 12345"
    client.prior_address = "456 Old St, Test City, TS 12345"
    client.gender = "Male"
    client.marital_status = "Single"
    client.date_of_birth = "1990-01-01"
    client.education_occupation = "College/Professional"
    client.drivers_license = "DL12345678"
    client.license_state = "TS"
    client.ssn = "123-45-6789"
    client.relation = "Self"
    client.referred_by = "Website"
    
    quote = MagicMock()
    quote.id = 1
    quote.client = client
    quote.effective_date = "2023-01-01"
    quote.has_auto = True
    quote.has_home = True
    quote.has_specialty = True
    
    # Add auto data with multiple vehicles and drivers
    auto_data = {
        "current_carrier": "Test Insurance",
        "months_with_carrier": "24",
        "current_limits": "100/300/100",
        "quoting_limits": "250/500/250",
        "expiration_date": "2023-12-31",
        "premium": "$1,200",
        "additional_notes": "Test notes for auto policy",
        "vehicles": [
            {
                "year": "2020",
                "make": "Test Make 1",
                "model": "Test Model 1",
                "vin": "TESTVINN1234567",
                "usage": "Commute",
                "mileage": "12000",
                "driver": "Test Client",
                "comprehensive": "$500",
                "collision": "$500",
                "glass": "Yes",
                "towing": "Yes",
                "rental_reimbursement": "Yes",
                "financing": "Yes",
                "gap": "No"
            },
            {
                "year": "2018",
                "make": "Test Make 2",
                "model": "Test Model 2",
                "vin": "TESTVINN7654321",
                "usage": "Pleasure",
                "mileage": "5000",
                "driver": "Additional Driver",
                "comprehensive": "$1000",
                "collision": "$1000",
                "glass": "No",
                "towing": "No",
                "rental_reimbursement": "No",
                "financing": "No",
                "gap": "No"
            }
        ],
        "drivers": [
            {
                "firstName": "Test",
                "lastName": "Client",
                "dateOfBirth": "1990-01-01",
                "licenseNumber": "DL12345678",
                "licenseState": "TS",
                "gender": "Male",
                "maritalStatus": "Single"
            },
            {
                "firstName": "Additional",
                "lastName": "Driver",
                "dateOfBirth": "1992-02-02",
                "licenseNumber": "DL87654321",
                "licenseState": "TS",
                "gender": "Female",
                "maritalStatus": "Married"
            }
        ],
        "additional_insureds": [
            {
                "name": "Additional Insured 1",
                "relation": "Spouse",
                "gender": "Female",
                "marital_status": "Married",
                "date_of_birth": "1992-02-02",
                "education_occupation": "College/Professional",
                "drivers_license": "DL87654321",
                "license_state": "TS",
                "ssn": "987-65-4321"
            }
        ]
    }
    
    # Add home data
    home_data = {
        "coverage_type": "Homeowner",
        "current_insurance_carrier": "Test Home Ins",
        "months_with_carrier": "36",
        "expiration_date": "2023-12-31",
        "usage": "Primary",
        "form_type": "HO3",
        "number_household_members": "3",
        "year_built": "2000",
        "stories_style": "2 Story",
        "garage": "2 Car Attached",
        "sq_ft_above_ground": "2500",
        "percentage_finished_basement": "50%",
        "walk_out_basement": "Yes",
        "full_bath": "2",
        "three_qtr_bath": "1",
        "half_bath": "1",
        "siding_type": "Vinyl",
        "fire_place": "1",
        "woodstove": "No",
        "attached_structures": "Deck",
        "detached_structures": "Shed",
        "sprinkled": "No",
        "miles_from_fd": "2",
        "responding_fd": "City FD",
        "fire_hydrant_distance": "500ft",
        "roof_year_replaced": "2015",
        "roof_type": "Asphalt Shingle",
        "deck_type": "Wood",
        "deck_size": "400sqft",
        "porch_type": "Covered",
        "porch_size": "100sqft",
        "pool": "No",
        "pool_depth": "",
        "pool_diving_board": "",
        "trampoline": "No",
        "fence_type": "Wood",
        "fence_height": "6ft",
        "heating_system_year": "2010",
        "heating_system_type": "Forced Air",
        "electrical_year": "2000",
        "electrical_type_amps": "200A",
        "alarm": "Yes",
        "plumbing_year": "2000",
        "plumbing_material_type": "Copper",
        "sump_pump_limit": "$5,000",
        "service_line_limit": "$10,000",
        "septic_sewer": "Sewer",
        "flood_insurance": "No",
        "reconstruction_cost": "$350,000",
        "personal_property_value": "$175,000",
        "scheduled_items_type": "Jewelry",
        "scheduled_items_value": "$10,000",
        "e_bikes_detail_type": "None",
        "e_bikes_value": "$0",
        "pets": "Dog",
        "business_type": "None",
        "biting_pets": "No",
        "deductible": "$1,000",
        "wind_hail": "1%",
        "mortgage": "Yes",
        "premium": "$1,500",
        "bankruptcy_foreclosure": "No",
        "umbrella_value": "$1,000,000",
        "umbrella_uninsured_underinsured": "Yes",
        "misc_rec_vehicles_toys": "No"
    }
    
    # Add specialty data
    specialty_data = {
        "items": [
            {
                "type_toy": "Motorcycle",
                "year": "2021",
                "make": "Test Make",
                "model": "Test Model",
                "vin": "TESTVIN12345678",
                "comprehensive_deductible": "$500",
                "collision_deductible": "$500",
                "total_hp": "150",
                "max_speed": "120mph",
                "cc_size": "1000cc",
                "market_value": "$15,000",
                "stored": "Garage"
            },
            {
                "type_toy": "ATV",
                "year": "2019",
                "make": "Test ATV",
                "model": "Test Model",
                "vin": "ATVIN12345678",
                "comprehensive_deductible": "$250",
                "collision_deductible": "$250",
                "total_hp": "50",
                "max_speed": "45mph",
                "cc_size": "450cc",
                "market_value": "$5,000",
                "stored": "Shed"
            }
        ],
        "additional_info": "Test additional information for specialty items"
    }
    
    quote.auto_data = auto_data
    quote.home_data = home_data
    quote.specialty_data = specialty_data
    
    return quote

@pytest.fixture
def setup_docx_generation():
    """Setup required components for DOCX generation testing"""
    # Create a directory for test templates if it doesn't exist
    template_dir = TEST_DIR / "templates"
    template_dir.mkdir(exist_ok=True)
    
    # Create a basic test template (if needed for actual tests)
    # This would be a simple .docx file with placeholders
    
    # Return useful test configuration
    return {
        "template_dir": template_dir,
        "output_dir": SAMPLE_OUTPUT_DIR,
        "placeholder_pattern": r'{{(.*?)}}'
    }

def test_auto_field_mapping():
    """Test auto field mapping to placeholders"""
    from backend.app.services.document_generator import _add_auto_data_to_dictionary
    
    # Test data
    data = {}
    auto_data = {
        "current_carrier": "Test Insurance",
        "months_with_carrier": "24",
        "current_limits": "100/300/100",
        "quoting_limits": "250/500/250",
        "expiration_date": "2023-12-31",
        "premium": "$1,200",
        "additional_notes": "Test notes for auto policy"
    }
    
    # Call the function
    _add_auto_data_to_dictionary(data, auto_data)
    
    # Verify mappings
    assert data["a-current-carrier"] == "Test Insurance"
    assert data["a-mos-current-carrier"] == "24"
    assert data["a-climits"] == "100/300/100"
    assert data["a-qlimits"] == "250/500/250"
    assert data["a-exp-dt"] == "12/31/2023"  # Should be formatted
    assert data["aprem"] == "$1,200"
    assert data["auto-additional-notes"] == "Test notes for auto policy"

def test_vehicle_field_mapping():
    """Test vehicle field mapping to placeholders"""
    from backend.app.services.document_generator import _add_auto_data_to_dictionary
    
    # Test data
    data = {}
    auto_data = {
        "vehicles": [
            {
                "year": "2020",
                "make": "Test Make",
                "model": "Test Model",
                "vin": "TESTVINN1234567",
                "usage": "Commute",
                "mileage": "12000",
                "driver": "Test Driver",
                "comprehensive": "$500",
                "collision": "$500",
                "glass": True,
                "towing": True,
                "rental_reimbursement": True,
                "financing": True,
                "gap": False
            }
        ]
    }
    
    # Call the function
    _add_auto_data_to_dictionary(data, auto_data)
    
    # Verify mappings
    assert data["v1yr"] == "2020"
    assert data["v1make"] == "Test Make"
    assert data["v1model"] == "Test Model"
    assert data["v1vin"] == "TESTVINN1234567"
    assert data["v1usage"] == "Commute"
    assert data["v1mi"] == "12000"
    assert data["v1-driver"] == "Test Driver"
    assert data["v1comp"] == "$500"
    assert data["v1coll"] == "$500"
    assert data["v1glass"] == "Yes"
    assert data["v1tow"] == "Yes"
    assert data["v1rr"] == "Yes"
    assert data["v1fin"] == "Yes"
    assert data["v1gap"] == "No"

def test_home_field_mapping():
    """Test home field mapping to placeholders"""
    from backend.app.services.document_generator import _add_home_data_to_dictionary
    
    # Test data
    data = {}
    home_data = {
        "coverage_type": "Homeowner",
        "roof_year_replaced": "2015",
        "heating_system_year": "2010",
        "electrical_year": "2000",
        "plumbing_year": "2000",
        "reconstruction_cost": "$350,000",
        "expiration_date": "2023-12-31"
    }
    
    # Call the function
    _add_home_data_to_dictionary(data, home_data)
    
    # Verify mappings
    assert data["hcovtype"] == "Homeowner"
    assert data["hroofyrrep"] == "2015"
    assert data["hheatsysyr"] == "2010"
    assert data["h-electricalyr"] == "2000"
    assert data["hplumbyr"] == "2000"
    assert data["hreconcost"] == "$350,000"
    assert data["h-expiration-date"] == "12/31/2023"  # Should be formatted

def test_specialty_field_mapping():
    """Test specialty field mapping to placeholders"""
    from backend.app.services.document_generator import _add_specialty_data_to_dictionary
    
    # Test data
    data = {}
    specialty_data = {
        "items": [
            {
                "type_toy": "Motorcycle",
                "year": "2021",
                "make": "Test Make",
                "model": "Test Model",
                "vin": "TESTVIN12345678",
                "comprehensive_deductible": "$500",
                "market_value": "$15,000",
            }
        ],
        "additional_info": "Test additional information"
    }
    
    # Call the function
    _add_specialty_data_to_dictionary(data, specialty_data)
    
    # Verify mappings
    assert data["sp1-type-toy"] == "Motorcycle"
    assert data["sp1yr"] == "2021"
    assert data["sp1make"] == "Test Make"
    assert data["sp1model"] == "Test Model"
    assert data["sp1vin"] == "TESTVIN12345678"
    assert data["sp1comp"] == "$500"
    assert data["sp1val"] == "$15,000"
    assert data["sp-additional-info"] == "Test additional information"

def test_additional_insured_field_mapping():
    """Test additional insured field mapping to placeholders"""
    from backend.app.services.document_generator import _add_additional_insured_data
    
    # Create test data
    data = {}
    quote = MagicMock()
    quote.auto_data = {
        "additional_insureds": [
            {
                "name": "Additional Insured 1",
                "relation": "Spouse",
                "gender": "Female",
                "marital_status": "Married",
                "date_of_birth": "1992-02-02",
                "education_occupation": "College/Professional",
                "drivers_license": "DL87654321",
                "license_state": "TS",
                "ssn": "987-65-4321"
            }
        ]
    }
    
    # Call the function
    _add_additional_insured_data(data, quote)
    
    # Verify mappings
    assert data["ai2name"] == "Additional Insured 1"
    assert data["ai2rel"] == "Spouse"
    assert data["ai2g"] == "Female"
    assert data["ai2ms"] == "Married"
    assert data["ai2dob"] == "02/02/1992"  # Should be formatted
    assert data["ai2edocc"] == "College/Professional"
    assert data["ai2dln"] == "DL87654321"
    assert data["ai2dls"] == "TS"
    assert data["ai2ssn"] == "987-65-4321"
    
def test_document_generation(mock_quote, setup_docx_generation):
    """Test end-to-end document generation with many fields"""
    from backend.app.services.document_generator import generate_quote_document, _create_data_dictionary
    
    # Create data dictionary from mock_quote
    data = _create_data_dictionary(mock_quote, mock_quote.client)
    
    # Verify data dictionary contains expected keys
    assert len(data) > 50, "Data dictionary should contain many placeholders"
    
    # Check some key fields from each section
    # Primary insured
    assert data["pniname"] == "Test Client"
    assert data["pniphone"] == "555-123-4567"
    
    # Auto data
    assert data["a-current-carrier"] == "Test Insurance"
    
    # Home data
    assert data["hroofyrrep"] == "2015"
    
    # Specialty data
    assert data["sp1-type-toy"] == "Motorcycle"
    
    # Additional insured
    assert data["ai2name"] == "Additional Insured 1"

def test_complete_field_coverage():
    """Test that all placeholders from placeholders.txt are covered in data dictionary"""
    from backend.app.services.document_generator import _create_data_dictionary
    
    # Create a mock quote with all possible fields
    mock_quote = MagicMock()
    mock_client = MagicMock()
    
    # Set up all possible fields with dummy values
    for attribute in dir(mock_client):
        if not attribute.startswith('_'):
            setattr(mock_client, attribute, f"Test {attribute}")
    
    mock_quote.client = mock_client
    mock_quote.id = 1
    mock_quote.effective_date = "2023-01-01"
    mock_quote.has_auto = True
    mock_quote.has_home = True
    mock_quote.has_specialty = True
    
    # Create comprehensive auto_data with all fields
    auto_data = {}
    for i in range(1, 9):  # 8 vehicles
        auto_data[f"v{i}yr"] = f"20{20+i}"
        auto_data[f"v{i}make"] = f"Make {i}"
        # Add all other vehicle fields...
    
    home_data = {}
    for key in ["hroofyrrep", "hheatsysyr", "helecamps", "hplumbyr"]:
        home_data[key.replace('-', '_')] = f"Test {key}"
    
    specialty_data = {"items": []}
    for i in range(1, 9):  # 8 specialty items
        item = {}
        for key in ["type_toy", "year", "make", "model", "vin"]:
            item[key] = f"Test {key} {i}"
        specialty_data["items"].append(item)
    
    mock_quote.auto_data = auto_data
    mock_quote.home_data = home_data
    mock_quote.specialty_data = specialty_data
    
    # Generate data dictionary
    data = _create_data_dictionary(mock_quote, mock_client)
    
    # Load placeholders from placeholders.txt
    with open("placeholders.txt", "r") as f:
        lines = f.readlines()
    
    placeholders = []
    for line in lines:
        if "{{" in line and "}}" in line:
            # Extract placeholder name from format like "field-name: {{placeholder}}"
            placeholder = line.split("{{")[1].split("}}")[0].strip()
            placeholders.append(placeholder)
    
    # Check if all placeholders are covered
    missing_placeholders = []
    for placeholder in placeholders:
        if placeholder not in data:
            missing_placeholders.append(placeholder)
    
    assert len(missing_placeholders) == 0, f"Missing placeholders: {missing_placeholders}"

def test_process_docx_placeholders():
    """Test placeholder replacement in DOCX file"""
    from backend.app.services.document_generator import _process_docx_placeholders
    from docx import Document
    
    # Create a test document with placeholders
    doc = Document()
    doc.add_paragraph("Client: {{pniname}}")
    doc.add_paragraph("Phone: {{pniphone}}")
    doc.add_paragraph("Vehicle: {{v1yr}} {{v1make}} {{v1model}}")
    
    # Create a table with placeholders
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Auto Carrier"
    table.cell(0, 1).text = "{{a-current-carrier}}"
    table.cell(1, 0).text = "Home Carrier"
    table.cell(1, 1).text = "{{h-current-insurance-carrier}}"
    
    # Create data dictionary
    data = {
        "pniname": "Test Client",
        "pniphone": "555-123-4567",
        "v1yr": "2020",
        "v1make": "Test Make",
        "v1model": "Test Model",
        "a-current-carrier": "Test Auto Insurance",
        "h-current-insurance-carrier": "Test Home Insurance"
    }
    
    # Process placeholders
    _process_docx_placeholders(doc, data)
    
    # Verify replacements in paragraphs
    assert doc.paragraphs[0].text == "Client: Test Client"
    assert doc.paragraphs[1].text == "Phone: 555-123-4567"
    assert doc.paragraphs[2].text == "Vehicle: 2020 Test Make Test Model"
    
    # Verify replacements in table
    assert table.cell(0, 1).text == "Test Auto Insurance"
    assert table.cell(1, 1).text == "Test Home Insurance"

if __name__ == "__main__":
    pytest.main(["-v", "test_docx_generation.py"]) 